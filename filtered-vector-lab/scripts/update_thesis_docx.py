#!/usr/bin/env python3
"""
将生成的图表插入到毕业论文docx的第三、四、五章中。
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = Path(__file__).resolve().parent.parent
FIGURES_DIR = PROJECT_ROOT / "results" / "figures"

INPUT_DOCX = Path("/root/.kimi/sessions/4533fc940d7b12d5b363cd8a8ebb4e68/9dee12f0-f52d-4942-aff3-c8e2dd17ddfb/uploads/面向新型数据密集型应用的索引技术研究_b987a7.docx")
OUTPUT_DOCX = PROJECT_ROOT / "docs" / "面向新型数据密集型应用的索引技术研究_更新版.docx"
OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)

# 图表插入映射：在包含特定文本的段落后插入图片
# 格式: (匹配文本, 图片文件名, 图片标题, 宽度英寸)
INSERTIONS = [
    # 第四章
    ("图4-1  HNSW实验数据：查询时间随参数变化", "fig_hnsw_m_sensitivity.png", "图4-1  HNSW M参数敏感性（MS MARCO, efSearch=128）", 5.5),
    ("表4-1  HNSW M参数敏感性（MS MARCO, efSearch=128）", None, None, 0),  # 表格占位，后续处理
    ("图4-2  SWIRL无效动作掩码示例与奖励函数定义", None, None, 0),  # 已有图，不替换
    ("图4-3  SIEVE的HNSW搜索成本模型与子模优化示意图", None, None, 0),  # 已有图
    ("图4-4  窗口搜索ThreeSplit算法与理论复杂度分析", None, None, 0),  # 已有图

    # 第五章
    ("图5-1  各算法在MS MARCO数据集上的Recall@K对比", "fig_recall_k_comparison.png", "图5-1  各算法在MS MARCO数据集上的Recall@K对比", 5.5),
    ("表5-1  MS MARCO数据集召回率对比（K∈{1,10,100}）", None, None, 0),
    ("图5-2  各算法的Recall@10与QPS权衡曲线（MS MARCO数据集）", "fig_recall_qps_tradeoff.png", "图5-2  各算法的Recall@10与QPS权衡曲线（MS MARCO数据集）", 5.5),
    ("图5-3  HNSW efSearch参数敏感性分析（MS MARCO, M=48）", "fig_hnsw_ef_sensitivity.png", "图5-3  HNSW efSearch参数敏感性分析（MS MARCO, M=48）", 5.5),
    ("图5-4  HNSW M参数敏感性分析（MS MARCO, efSearch=128）", "fig_hnsw_m_sensitivity.png", "图5-4  HNSW M参数敏感性分析（MS MARCO, efSearch=128）", 5.5),
    ("图5-5  消融实验：各组件的独立贡献分析", "fig_ablation.png", "图5-5  消融实验：各组件的独立贡献分析", 5.5),
    ("表5-2  消融实验结果（MS MARCO数据集）", None, None, 0),
    ("图5-6  各算法在三个数据集上的Recall@10对比", "fig_cross_dataset.png", "图5-6  各算法在三个数据集上的Recall@10对比", 5.5),
    ("图5-7  不同过滤强度下的Recall@10与延迟对比（Enron Email）", "fig_filter_query.png", "图5-7  不同过滤强度下的Recall@10与延迟对比（Enron Email）", 5.5),
    ("表5-3  过滤查询性能对比（Recall@10/延迟）", None, None, 0),
]


def insert_image_after_paragraph(doc, match_text, img_path, caption, width_inches):
    """在包含匹配文本的段落后插入图片和标题。"""
    for i, para in enumerate(doc.paragraphs):
        if match_text in para.text:
            print(f"  Found: '{match_text[:40]}...' at para {i}")
            # 在该段落后插入图片
            new_para = doc.paragraphs[i]._element
            new_p = doc.add_paragraph()
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = new_p.add_run()
            run.add_picture(str(img_path), width=Inches(width_inches))
            # 移动新段落到目标位置之后
            new_p._element.getparent().insert(
                doc.paragraphs[i]._element.getparent().index(doc.paragraphs[i]._element) + 1,
                new_p._element
            )

            # 插入标题段落
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_p.add_run(caption)
            cap_run.font.size = Pt(10.5)
            cap_run.font.name = "Times New Roman"
            cap_p._element.getparent().insert(
                new_p._element.getparent().index(new_p._element) + 1,
                cap_p._element
            )
            return True
    print(f"  NOT FOUND: '{match_text[:40]}...'")
    return False


def main():
    print(f"Loading document: {INPUT_DOCX}")
    doc = Document(str(INPUT_DOCX))
    print(f"Total paragraphs: {len(doc.paragraphs)}")

    inserted = 0
    for match_text, img_file, caption, width in INSERTIONS:
        if img_file is None:
            continue
        img_path = FIGURES_DIR / img_file
        if not img_path.exists():
            print(f"[SKIP] Image not found: {img_path}")
            continue
        if insert_image_after_paragraph(doc, match_text, img_path, caption, width):
            inserted += 1

    print(f"\nInserted {inserted} images.")
    doc.save(str(OUTPUT_DOCX))
    print(f"Saved updated document to: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
